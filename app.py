from __future__ import annotations

import html
import io
import json
import re
import sys
from dataclasses import dataclass
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import urlparse

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - surfaced in the UI at runtime
    PdfReader = None

try:
    from docx import Document
except Exception:  # pragma: no cover - surfaced in the UI at runtime
    Document = None


ROOT = Path(__file__).resolve().parent
STATIC_DIR = ROOT / "static"
MAX_UPLOAD_BYTES = 12 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx"}
MAX_PASTE_CHARACTERS = 90_000


@dataclass
class StudySection:
    title: str
    body: str
    bullets: list[str]


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(data: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("缺少 pypdf，無法解析 PDF。")

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"第 {index} 頁\n{page_text.strip()}")
    return "\n\n".join(pages)


def extract_text_from_docx(data: bytes) -> str:
    if Document is None:
        raise RuntimeError("缺少 python-docx，無法解析 DOCX。")

    document = Document(io.BytesIO(data))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            blocks.append(value)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(blocks)


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("目前支援 txt、md、html、pdf、docx。")

    if suffix == ".pdf":
        return extract_text_from_pdf(data)
    if suffix == ".docx":
        return extract_text_from_docx(data)

    for encoding in ("utf-8-sig", "utf-8", "big5", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def split_into_sections(text: str) -> list[StudySection]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if not paragraphs:
        return []

    sections: list[StudySection] = []
    current_title = "重點整理"
    current_body: list[str] = []

    def flush() -> None:
        nonlocal current_body
        body = "\n\n".join(current_body).strip()
        if body:
            sections.append(
                StudySection(
                    title=current_title,
                    body=body,
                    bullets=make_bullets(body),
                )
            )
        current_body = []

    for paragraph in paragraphs:
        maybe_heading = paragraph.replace("#", "").strip()
        is_heading = (
            len(maybe_heading) <= 42
            and "\n" not in maybe_heading
            and (
                paragraph.startswith("#")
                or re.match(r"^(\d+[\).、]|第.+[章節]|[一二三四五六七八九十]+、)", maybe_heading)
            )
        )

        if is_heading:
            flush()
            current_title = maybe_heading
        else:
            current_body.append(paragraph)

    flush()
    if not sections:
        sections.append(StudySection(title="重點整理", body=text, bullets=make_bullets(text)))

    if len(sections) == 1 and len(paragraphs) > 4:
        sections = regroup_paragraphs(paragraphs)

    return sections[:5]


def regroup_paragraphs(paragraphs: list[str]) -> list[StudySection]:
    group_size = max(1, round(len(paragraphs) / 4))
    grouped: list[StudySection] = []
    for index in range(0, len(paragraphs), group_size):
        body_parts = paragraphs[index : index + group_size]
        body = "\n\n".join(body_parts)
        title = infer_title(body, len(grouped) + 1)
        grouped.append(StudySection(title=title, body=body, bullets=make_bullets(body)))
    return grouped[:5]


def split_sentences(text: str) -> list[str]:
    clean = re.sub(r"\s+", " ", text).strip()
    parts = re.split(r"(?<=[。！？.!?])\s*|(?<=；)\s*", clean)
    return [part.strip(" -•\t") for part in parts if len(part.strip()) >= 12]


def make_bullets(text: str) -> list[str]:
    sentences = split_sentences(text)
    if not sentences:
        lines = [line.strip(" -•\t") for line in text.splitlines() if len(line.strip()) >= 12]
        sentences = lines

    scored = sorted(
        sentences,
        key=lambda item: (
            contains_keyword(item),
            min(len(item), 120),
        ),
        reverse=True,
    )
    return [compress_sentence(sentence) for sentence in scored[:4]]


def compress_sentence(sentence: str) -> str:
    sentence = re.sub(r"\s+", " ", sentence).strip(" -•\t")
    replacements = {
        "然後": " -> ",
        "而且": "；",
        "所以": " -> ",
        "因此": " -> ",
        "導致": " -> ",
        "造成": " -> ",
        "because": " -> ",
        "therefore": " -> ",
    }
    for source, target in replacements.items():
        sentence = sentence.replace(source, target)
    sentence = re.sub(r"(其實|就是|非常|相當|很多很多|基本上)", "", sentence)
    sentence = re.sub(r"\s+", " ", sentence)
    return sentence[:120].rstrip("，,。；;") + ("…" if len(sentence) > 120 else "")


def infer_title(text: str, fallback_index: int) -> str:
    sentences = split_sentences(text)
    first = sentences[0] if sentences else text.splitlines()[0]
    first = re.sub(r"^[#\d\s).、]+", "", first).strip()
    first = re.split(r"[：:。.!?！？]", first)[0].strip()
    return first[:22] or f"重點 {fallback_index}"


def focus_terms(text: str) -> list[str]:
    cjk_terms = re.findall(r"[\u4e00-\u9fff]{2,8}", text)
    latin_terms = re.findall(r"[A-Za-z][A-Za-z0-9_-]{4,}", text)
    stop_terms = {
        "這個",
        "因此",
        "所以",
        "可以",
        "使用",
        "內容",
        "文字",
        "資料",
        "文件",
        "進行",
        "需要",
        "整理",
        "確認",
    }
    counts: dict[str, int] = {}
    for term in cjk_terms + latin_terms:
        if term in stop_terms:
            continue
        counts[term] = counts.get(term, 0) + 1

    ranked = sorted(counts, key=lambda item: (counts[item], len(item)), reverse=True)
    return ranked[:3]


def render_inline(text: str, terms: list[str]) -> str:
    rendered = html.escape(text)
    for term in terms:
        escaped = html.escape(term)
        rendered = re.sub(
            re.escape(escaped),
            f"<strong>{escaped}</strong>",
            rendered,
            count=1,
        )
    return rendered


def render_bullets(bullets: list[str], terms: list[str]) -> str:
    return "\n".join(f"<li>{render_inline(bullet, terms)}</li>" for bullet in bullets)


def render_paragraphs(body: str, terms: list[str]) -> str:
    paragraphs = []
    for paragraph in re.split(r"\n\s*\n", body):
        clean = paragraph.strip()
        if not clean:
            continue
        if len(clean) > 220:
            for sentence in split_sentences(clean)[:4]:
                paragraphs.append(f"<p>{render_inline(compress_sentence(sentence), terms)}</p>")
        else:
            paragraphs.append(f"<p>{render_inline(compress_sentence(clean), terms)}</p>")
    return "\n".join(paragraphs)


def make_comparison_table(section: StudySection, terms: list[str]) -> str:
    comparison_words = ("優點", "缺點", "風險", "比較", "差異", "策略", "方法", "方案", "適用")
    if not any(word in section.body for word in comparison_words):
        return ""

    bullets = section.bullets[:3]
    if len(bullets) < 2:
        bullets = [compress_sentence(sentence) for sentence in split_sentences(section.body)[:3]]
    if len(bullets) < 2:
        return ""

    rows = []
    for index, bullet in enumerate(bullets, start=1):
        rows.append(
            "<tr>"
            f"<td>{render_inline(f'觀點 {index}', terms)}</td>"
            f"<td>{render_inline(bullet, terms)}</td>"
            f"<td>{render_inline('需搭配原文脈絡判讀', terms)}</td>"
            "</tr>"
        )
    return f"""
      <div class="table-wrap">
        <table>
          <thead>
            <tr><th>比較項目</th><th>重點訊號</th><th>注意事項</th></tr>
          </thead>
          <tbody>{''.join(rows)}</tbody>
        </table>
      </div>
    """


def contains_keyword(sentence: str) -> int:
    keywords = [
        "重要",
        "因此",
        "所以",
        "目的",
        "方法",
        "結果",
        "優點",
        "缺點",
        "比較",
        "definition",
        "method",
        "result",
        "important",
        "because",
    ]
    lowered = sentence.lower()
    return sum(1 for keyword in keywords if keyword.lower() in lowered)


def estimate_reading_minutes(text: str) -> int:
    cjk_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    latin_words = len(re.findall(r"[A-Za-z0-9_]+", text))
    minutes = round((cjk_chars / 450) + (latin_words / 220))
    return max(1, minutes)


def make_study_html(title: str, source_name: str, text: str) -> str:
    sections = split_into_sections(text)
    escaped_title = html.escape(title)
    escaped_source = html.escape(source_name)
    reading_minutes = estimate_reading_minutes(text)
    word_count = len(re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9_]+", text))

    nav = "\n".join(
        f'<a href="#section-{index}">{html.escape(section.title)}</a>'
        for index, section in enumerate(sections, start=1)
    )

    section_html = []
    for index, section in enumerate(sections, start=1):
        terms = focus_terms(section.body)
        bullets = render_bullets(section.bullets, terms)
        paragraphs = render_paragraphs(section.body, terms)
        insight = section.bullets[0] if section.bullets else infer_title(section.body, index)
        table = make_comparison_table(section, terms)
        section_html.append(
            f"""
            <section id="section-{index}" class="study-section">
              <h2>{html.escape(section.title)}</h2>
              <blockquote>
                <strong>核心結論</strong>
                <p>{render_inline(insight, terms)}</p>
              </blockquote>
              <div class="takeaway">
                <strong>整理後重點</strong>
                <ul>{bullets}</ul>
              </div>
              {table}
              <div class="content-flow">{paragraphs}</div>
            </section>
            """
        )

    return f"""<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{escaped_title}</title>
  <style>
    :root {{
      color-scheme: light;
      --ink: #17202a;
      --muted: #667085;
      --paper: #ffffff;
      --line: #e8e8ed;
      --accent: #0066cc;
      --accent-soft: #f5f5f7;
      --mark: #0071e3;
      --quote: #f2f7ff;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background: #f5f5f7;
      color: var(--ink);
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Noto Sans TC", sans-serif;
      line-height: 1.75;
    }}
    .page {{
      max-width: 980px;
      margin: 0 auto;
      padding: 34px 20px 60px;
    }}
    header {{
      padding: 24px 0 18px;
      text-align: center;
    }}
    h1 {{
      margin: 0 0 10px;
      font-size: clamp(1.75rem, 3.5vw, 3rem);
      line-height: 1;
      letter-spacing: 0;
      font-weight: 800;
    }}
    .meta {{
      display: flex;
      flex-wrap: wrap;
      justify-content: center;
      gap: 10px;
      color: var(--muted);
    }}
    .meta span {{
      background: #ffffff;
      border: 1px solid var(--line);
      border-radius: 999px;
      padding: 5px 12px;
    }}
    nav {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      padding: 20px 0;
      justify-content: center;
    }}
    nav a {{
      color: var(--accent);
      background: #ffffff;
      border: 1px solid var(--line);
      text-decoration: none;
      border-radius: 999px;
      padding: 7px 12px;
      font-weight: 700;
    }}
    .study-section {{
      background: var(--paper);
      border: 1px solid var(--line);
      border-radius: 18px;
      padding: clamp(22px, 4vw, 44px);
      margin-top: 14px;
    }}
    blockquote {{
      margin: 0 0 18px;
      padding: 16px 20px;
      border-left: 4px solid var(--mark);
      border-radius: 14px;
      background: var(--quote);
    }}
    blockquote p {{
      margin: 6px 0 0;
    }}
    h2 {{
      margin: 0 0 16px;
      font-size: clamp(1.35rem, 3vw, 2rem);
      line-height: 1.25;
      letter-spacing: 0;
    }}
    .takeaway {{
      background: var(--accent-soft);
      border-radius: 14px;
      padding: 16px 20px;
      margin-bottom: 22px;
    }}
    strong {{
      font-weight: 800;
    }}
    .table-wrap {{
      overflow-x: auto;
      margin: 18px 0 22px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 14px;
      overflow: hidden;
    }}
    th, td {{
      padding: 12px 14px;
      border-bottom: 1px solid var(--line);
      text-align: left;
      vertical-align: top;
    }}
    th {{
      background: var(--accent-soft);
      font-size: .92rem;
    }}
    .takeaway ul {{
      margin: 8px 0 0;
      padding-left: 22px;
    }}
    .content-flow p {{
      margin: 0 0 14px;
      text-align: justify;
    }}
    @media print {{
      body {{ background: white; }}
      .page {{ padding: 0; }}
    }}
  </style>
</head>
<body>
  <main class="page">
    <header>
      <h1>{escaped_title}</h1>
      <div class="meta">
        <span>來源：{escaped_source}</span>
        <span>約 {reading_minutes} 分鐘閱讀</span>
        <span>{word_count} 個字詞單位</span>
      </div>
    </header>
    <nav aria-label="章節導覽">{nav}</nav>
    {"".join(section_html)}
  </main>
</body>
</html>"""


def parse_multipart(body: bytes, content_type: str) -> tuple[str, bytes]:
    match = re.search(r"boundary=(.+)", content_type)
    if not match:
        raise ValueError("缺少 multipart boundary。")

    boundary = ("--" + match.group(1).strip().strip('"')).encode()
    for part in body.split(boundary):
        if b"Content-Disposition" not in part:
            continue
        header, _, payload = part.partition(b"\r\n\r\n")
        name_match = re.search(rb'filename="([^"]+)"', header)
        if not name_match:
            continue
        filename = name_match.group(1).decode("utf-8", errors="replace")
        return filename, payload.rstrip(b"\r\n-")
    raise ValueError("找不到上傳檔案。")


def json_response(handler: SimpleHTTPRequestHandler, status: HTTPStatus, payload: dict) -> None:
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def build_result(filename: str, title: str, text: str) -> dict:
    clean_text = normalize_text(text)
    if not clean_text:
        raise ValueError("沒有可整理的文字。")

    clean_title = title.strip() or Path(filename).stem.replace("_", " ").replace("-", " ").strip() or "學習筆記"
    sections = split_into_sections(clean_text)
    return {
        "filename": filename,
        "title": clean_title,
        "plainText": clean_text,
        "html": make_study_html(clean_title, filename, clean_text),
        "stats": {
            "characters": len(clean_text),
            "sections": len(sections),
            "readingMinutes": estimate_reading_minutes(clean_text),
        },
    }


class StudyHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(STATIC_DIR), **kwargs)

    def do_POST(self) -> None:
        path = urlparse(self.path).path
        if path not in {"/api/convert", "/api/paste"}:
            json_response(self, HTTPStatus.NOT_FOUND, {"error": "找不到 API。"})
            return

        try:
            length = int(self.headers.get("Content-Length", "0"))
            if length > MAX_UPLOAD_BYTES:
                raise ValueError("檔案太大，請使用 12MB 以下的文件。")

            body = self.rfile.read(length)
            if path == "/api/paste":
                payload = json.loads(body.decode("utf-8"))
                pasted_text = str(payload.get("text", ""))
                if len(pasted_text) > MAX_PASTE_CHARACTERS:
                    raise ValueError("貼上的文字太長，請控制在 90,000 字元以下。")
                title = str(payload.get("title", "")).strip() or "貼上文字筆記"
                json_response(self, HTTPStatus.OK, build_result("pasted-text.txt", title, pasted_text))
                return

            filename, data = parse_multipart(body, self.headers.get("Content-Type", ""))
            text = normalize_text(extract_text(filename, data))
            if not text:
                raise ValueError("沒有抽取到文字。掃描圖像型 PDF 需要先 OCR。")

            title = Path(filename).stem.replace("_", " ").replace("-", " ").strip() or "學習筆記"
            json_response(self, HTTPStatus.OK, build_result(filename, title, text))
        except Exception as exc:
            json_response(self, HTTPStatus.BAD_REQUEST, {"error": str(exc)})


def run(port: int = 5177) -> None:
    server = ThreadingHTTPServer(("127.0.0.1", port), StudyHandler)
    print(f"Study HTML Maker running at http://127.0.0.1:{port}")
    server.serve_forever()


if __name__ == "__main__":
    selected_port = int(sys.argv[1]) if len(sys.argv) > 1 else 5177
    run(selected_port)
